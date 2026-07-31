from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BLACK = RGBColor(0, 0, 0)

def set_black(run):
    run.font.color.rgb = BLACK

def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_black(run)
        run.font.color.rgb = BLACK
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    set_black(run)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.size = Pt(11)
        set_black(r1)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        set_black(r2)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        set_black(run)
    return p

def add_table(doc, rows_data, header=True):
    n_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=n_cols)
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(11)
            set_black(run)
            if i == 0 and header:
                run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return table

doc = Document()

# --- Márgenes ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

# Forzar estilos de heading en negro
for i in range(1, 5):
    style_name = f'Heading {i}'
    try:
        style = doc.styles[style_name]
        style.font.color.rgb = BLACK
    except Exception:
        pass

# ===== TÍTULO =====
title = doc.add_heading('Propuesta de Servicios Profesionales', 0)
for run in title.runs:
    run.font.color.rgb = BLACK
    run.font.size = Pt(22)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_heading('Rediseño Web + Email Corporativo — Wilson Hotel', 2)
for run in sub.runs:
    run.font.color.rgb = BLACK
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Tabla de encabezado
header_data = [
    ['Para:', 'Wilson Hotel — Alvarado 950, Salta'],
    ['De:', 'Agustín Pedernera'],
    ['Fecha:', '20 de junio de 2026'],
    ['Vigencia:', '30 días'],
]
t = doc.add_table(rows=4, cols=2)
t.style = 'Table Grid'
for i, (label, value) in enumerate(header_data):
    row = t.rows[i]
    c0 = row.cells[0]
    c0.text = ''
    r = c0.paragraphs[0].add_run(label)
    r.font.bold = True
    r.font.size = Pt(11)
    set_black(r)
    c1 = row.cells[1]
    c1.text = ''
    r2 = c1.paragraphs[0].add_run(value)
    r2.font.size = Pt(11)
    set_black(r2)

doc.add_paragraph()

# ===== RESUMEN =====
heading(doc, 'Resumen', 2)
body(doc,
    'Desarrollo de un nuevo sitio web institucional, personalización visual del motor de reservas Winpax '
    'y regularización del dominio wilsonhotel.com.ar (que vence el 15/07/2026).\n\n'
    'Esta versión contempla el escenario en que el webmaster actual transfiere las cuentas de hosting y email al hotel. '
    'En ese caso, se continúa operando sobre la infraestructura existente en DonWeb sin necesidad de migración.')

# ===== SITUACIÓN ACTUAL =====
heading(doc, 'Situación Actual', 2)
bullet(doc, 'sobre WordPress con PHP 7.4 (sin soporte desde 2022), diseño desactualizado, sin optimización móvil, gestionado por un webmaster externo.', bold_prefix='Sitio web ')
bullet(doc, 'en servidor con configuración de seguridad deficiente: sin firma digital (DKIM), sin protección activa contra suplantación (DMARC pasivo).', bold_prefix='Email ')
bullet(doc, 'wilsonhotel.com.ar vence el 15 de julio de 2026 — renovar es prerequisito absoluto.', bold_prefix='Dominio ')

# ===== ALCANCE =====
heading(doc, 'Alcance', 2)

heading(doc, '1. Sitio Web Nuevo', 3)
body(doc, 'Landing page institucional alineada al brandbook del hotel, responsive, orientada a generar reservas directas.')
bullet(doc, 'Diseño UI/UX con identidad de marca (paleta, tipografías, sistema de logos)')
bullet(doc, 'Secciones: Hero, Habitaciones, Servicios, Ubicación, Galería, Salones, Contacto')
bullet(doc, 'Integración con motor Winpax (botón de reserva)')
bullet(doc, 'Botón flotante de WhatsApp')
bullet(doc, 'SEO básico (meta tags, Open Graph, sitemap)')
bullet(doc, 'Performance optimizada (objetivo >90 en PageSpeed)')
bullet(doc, 'Deploy en hosting DonWeb (cuenta transferida del webmaster)')

heading(doc, '2. Personalización del Motor de Reservas (Winpax)', 3)
body(doc, 'Coordinación directa con el equipo de Winpax para adaptar la apariencia visual del motor de reservas a la estética del nuevo sitio y del manual de marca.')
bullet(doc, 'Definición de guía de estilo para el motor: colores, tipografías, logo')
bullet(doc, 'Comunicación con Winpax para implementar los cambios visuales')
bullet(doc, 'Seguimiento hasta lograr coherencia visual entre el sitio y el motor')

heading(doc, '3. Dominio', 3)
bullet(doc, 'Renovación urgente del dominio en NIC Argentina')
bullet(doc, 'Coordinación con webmaster para transferencia de la cuenta de hosting y email')
bullet(doc, 'Ajuste de configuración DNS si fuera necesario')
bullet(doc, 'El hotel recibe credenciales propias de todas las cuentas')

heading(doc, '4. Puesta en Marcha', 3)
bullet(doc, 'Testing cross-browser y móvil')
bullet(doc, 'Capacitación básica + documentación de accesos')

# ===== ACCESOS QUE DEBE PROVEER EL WEBMASTER / HOTEL =====
heading(doc, 'Accesos que el webmaster / hotel debe proveer', 2)
body(doc,
    'Para iniciar el proyecto, necesito que el hotel obtenga o solicite al webmaster actual los siguientes accesos. '
    'Estos son prerequisitos: sin ellos no es posible renovar el dominio, desplegar el sitio ni configurar el email.')

heading(doc, 'Del webmaster actual', 3)
bullet(doc, 'Usuario y contraseña del panel de control de DonWeb (cPanel o panel propio) — para acceder al hosting donde vive el sitio actual')
bullet(doc, 'Credenciales FTP o acceso de archivos al servidor')
bullet(doc, 'Acceso al panel de administración de WordPress actual (usuario administrador)')
bullet(doc, 'Acceso a la cuenta de DonWeb Email — panel de gestión de casillas de correo existentes')
bullet(doc, 'Confirmación escrita de que la transferencia de las cuentas DonWeb al hotel está autorizada')

heading(doc, 'Del hotel (directamente)', 3)
bullet(doc, 'Acceso a la cuenta NIC Argentina asociada al dominio wilsonhotel.com.ar — para renovar antes del 15/07/2026')
bullet(doc, 'Credenciales o contacto directo con Winpax (nombre de usuario del hotel en el motor de reservas) — para coordinar la personalización visual')
bullet(doc, 'Persona de contacto con autoridad para aprobar decisiones de diseño y contenido')
bullet(doc, 'Imágenes y contenido que el hotel quiera usar en el sitio nuevo (fotos de habitaciones, salones, exteriores)')

body(doc, 'Una vez que el webmaster transfiera las cuentas DonWeb, el hotel pasa a ser titular directo de hosting y email, sin depender de intermediarios.', bold=False)

# ===== CRONOGRAMA =====
heading(doc, 'Cronograma', 2)
add_table(doc, [
    ['Semana', 'Actividad'],
    ['1', 'Renovación dominio + transferencia de cuentas desde webmaster'],
    ['2', 'Deploy sitio web nuevo + configuración DNS'],
    ['3', 'Coordinación visual con Winpax + SEO + testing'],
    ['4', 'Puesta en marcha'],
])
doc.add_paragraph()
body(doc, 'Prerequisitos para iniciar: acceso a NIC Argentina, que el webmaster confirme la transferencia, persona de contacto para decisiones.', bold=True)

# ===== INVERSIÓN =====
heading(doc, 'Inversión', 2)

heading(doc, 'Proyecto (pago único)', 3)
add_table(doc, [
    ['Concepto', 'Inversión'],
    ['Sitio web nuevo (diseño, desarrollo, despliegue)', 'ARS $100.000'],
    ['SEO básico + performance', 'ARS $50.000'],
    ['Total', 'ARS $150.000'],
])

doc.add_paragraph()
heading(doc, 'Costos fijos (paga el hotel a los proveedores)', 3)
add_table(doc, [
    ['Servicio', 'Costo'],
    ['DonWeb Hosting (cuenta transferida, plan actual)', 'Según plan vigente'],
    ['DonWeb Email (cuenta transferida, plan actual)', 'Según plan vigente'],
    ['Dominio .com.ar (NIC Argentina)', '~ARS 8.500/año'],
])

doc.add_paragraph()
heading(doc, 'Mantenimiento mensual', 3)
add_table(doc, [
    ['Incluye', 'Inversión'],
    ['Soporte técnico, actualizaciones, monitoreo, ajustes, backups', 'ARS $40.000/mes'],
])

# ===== CONDICIONES =====
doc.add_paragraph()
heading(doc, 'Condiciones', 2)
bullet(doc, '50% al aprobar / 50% contra entrega', bold_prefix='Pago: ')
bullet(doc, 'El código y las credenciales quedan a nombre del hotel al completar el pago', bold_prefix='Propiedad: ')
bullet(doc, 'Se comienza a pagar el mes siguiente a la implementación', bold_prefix='Mantenimiento: ')
bullet(doc, 'La información del relevamiento se trata como confidencial', bold_prefix='Confidencialidad: ')

# ===== PRÓXIMOS PASOS =====
heading(doc, 'Próximos Pasos', 2)
steps = [
    'Aprobación de esta propuesta',
    'Pago del anticipo (50%)',
    'Coordinar con webmaster la transferencia de cuentas',
    'Entrega de accesos (dominio, hosting, email, Winpax)',
    'Inicio — primera acción: renovar dominio',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(step)
    run.font.size = Pt(11)
    set_black(run)

doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run('Agustín Pedernera — agustinpedernera147@gmail.com — Junio 2026')
r.font.size = Pt(10)
r.font.italic = True
set_black(r)

out = '/Users/147pop/Documents/GitHub/hotelwilson/entrevista/propuesta-wilson-hotel-v2.docx'
doc.save(out)
print(f'Guardado: {out}')
